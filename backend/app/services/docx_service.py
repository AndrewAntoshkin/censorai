import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.schemas.analysis import AnalysisResponse, AnalysisSummary


RISK_LABELS: dict[str, str] = {
    "drugs": "Наркотики",
    "weapons": "Оружие",
    "violence": "Насилие",
    "sexual_content": "Сексуальный контент",
    "profanity": "Нецензурная лексика",
    "illegal_actions": "Незаконные действия",
    "alcohol": "Алкоголь",
    "smoking": "Курение",
    "animal_cruelty": "Жестокое обращение с животными",
    "forbidden_symbols": "Запрещённая символика",
    "text_in_frame": "Текст в кадре",
    "discreditation_values": "Дискредитация ценностей",
    "propaganda": "Пропаганда",
    "crime_glorification": "Героизация преступлений",
    "excessive_cruelty": "Чрезмерная жестокость",
}

RISK_LEVEL_LABELS: dict[str, str] = {
    "critical": "Критический",
    "warning": "Предупреждение",
    "info": "Информация",
}

RECOMMENDATION_LABELS: dict[str, str] = {
    "remove": "Удалить",
    "shorten": "Сократить",
    "mute": "Заглушить звук",
    "blur": "Заблюрить",
}


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text) if text else "—")
    run.bold = bold
    run.font.size = Pt(size)


def generate_report(analysis: AnalysisResponse) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # --- Title page ---
    title = doc.add_heading("фреймчек — Отчёт анализа", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if analysis.video_title:
        run = info_para.add_run(f"Файл: {analysis.video_title}\n")
        run.font.size = Pt(14)
    if analysis.duration:
        run = info_para.add_run(f"Длительность: {analysis.duration}\n")
        run.font.size = Pt(12)
    date_str = analysis.analyzed_at.strftime("%d.%m.%Y %H:%M") if analysis.analyzed_at else datetime.now().strftime("%d.%m.%Y %H:%M")
    run = info_para.add_run(f"Дата анализа: {date_str}")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Summary ---
    doc.add_heading("Сводка", level=1)

    summary = _build_summary(analysis)

    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Light Grid Accent 1"

    rows_data = [
        ("Всего сцен", str(summary.total_scenes)),
        ("Сцен с рисками", str(summary.risky_scenes)),
        ("Критических", str(summary.critical_count)),
        ("Предупреждений", str(summary.warning_count)),
    ]
    for i, (label, value) in enumerate(rows_data):
        _set_cell_text(summary_table.rows[i].cells[0], label, bold=True)
        _set_cell_text(summary_table.rows[i].cells[1], value)

    if summary.risk_categories:
        doc.add_paragraph("")
        doc.add_heading("Обнаруженные категории рисков", level=2)
        for cat, count in summary.risk_categories.items():
            label = RISK_LABELS.get(cat, cat)
            doc.add_paragraph(f"• {label}: {count}", style="List Bullet")

    doc.add_page_break()

    # --- Full scene table ---
    doc.add_heading("Разбивка по сценам", level=1)

    if analysis.scenes:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["№", "Время", "Описание", "Риск", "Уровень"]
        for i, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], header, bold=True, size=10)

        for scene in analysis.scenes:
            row = table.add_row()
            _set_cell_text(row.cells[0], str(scene.scene_number))
            time_range = f"{scene.start_time or '?'} — {scene.end_time or '?'}"
            _set_cell_text(row.cells[1], time_range)
            _set_cell_text(row.cells[2], scene.description or "—")
            risk_label = RISK_LABELS.get(scene.risk, scene.risk) if scene.risk else "—"
            _set_cell_text(row.cells[3], risk_label)
            level_label = RISK_LEVEL_LABELS.get(scene.risk_level, scene.risk_level) if scene.risk_level else "—"
            _set_cell_text(row.cells[4], level_label)

            if scene.risk_level == "critical":
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # --- Risky scenes detail ---
    risky_scenes = [s for s in analysis.scenes if s.risk]
    if risky_scenes:
        doc.add_heading("Детали рисковых сцен", level=1)
        for scene in risky_scenes:
            doc.add_heading(
                f"Сцена {scene.scene_number} ({scene.start_time} — {scene.end_time})",
                level=2,
            )
            if scene.description:
                doc.add_paragraph(scene.description)

            detail_table = doc.add_table(rows=0, cols=2)
            detail_table.style = "Light Grid Accent 1"

            details = [
                ("Категория риска", RISK_LABELS.get(scene.risk, scene.risk)),
                ("Уровень", RISK_LEVEL_LABELS.get(scene.risk_level, scene.risk_level)),
                ("Вероятность", f"{scene.probability:.0%}" if scene.probability is not None else "—"),
                ("Причина", scene.reason),
                ("Цитата", scene.quote),
                ("Текст в кадре", scene.text_in_frame),
                ("Рекомендация", RECOMMENDATION_LABELS.get(scene.recommendation, scene.recommendation) if scene.recommendation else "—"),
            ]
            for label, value in details:
                row = detail_table.add_row()
                _set_cell_text(row.cells[0], label, bold=True)
                _set_cell_text(row.cells[1], value or "—")

            doc.add_paragraph("")

    # --- Recommendations ---
    doc.add_heading("Рекомендации", level=1)
    if risky_scenes:
        seen_recs: set[str] = set()
        for scene in risky_scenes:
            if scene.recommendation and scene.recommendation not in seen_recs:
                seen_recs.add(scene.recommendation)
                rec_label = RECOMMENDATION_LABELS.get(scene.recommendation, scene.recommendation)
                scenes_with_rec = [
                    str(s.scene_number) for s in risky_scenes if s.recommendation == scene.recommendation
                ]
                doc.add_paragraph(
                    f"• {rec_label} — сцены: {', '.join(scenes_with_rec)}",
                    style="List Bullet",
                )
    else:
        doc.add_paragraph("Рисков не обнаружено. Видео соответствует требованиям.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_summary(analysis: AnalysisResponse) -> AnalysisSummary:
    total = len(analysis.scenes)
    risky = [s for s in analysis.scenes if s.risk]
    categories: dict[str, int] = {}
    critical = 0
    warning = 0

    for scene in risky:
        if scene.risk:
            categories[scene.risk] = categories.get(scene.risk, 0) + 1
        if scene.risk_level == "critical":
            critical += 1
        elif scene.risk_level == "warning":
            warning += 1

    return AnalysisSummary(
        total_scenes=total,
        risky_scenes=len(risky),
        risk_categories=categories,
        critical_count=critical,
        warning_count=warning,
    )
